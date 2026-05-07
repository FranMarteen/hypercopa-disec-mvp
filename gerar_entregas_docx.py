"""Gera docs/ENTREGA_TIME1_MODELOS.docx e docs/ENTREGA_TIME2_AGENTE.docx
a partir dos markdowns correspondentes, com identidade visual BB.

Conversor enxuto de Markdown -> python-docx, suficiente para os documentos
de entrega (cabeçalhos, parágrafos, tabelas, listas, código inline, blockquote).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent

AMARELO_HEX = "FAE128"
AZUL_HEX = "003DA5"
AZUL_ESCURO_HEX = "002D72"
CINZA_FUNDO_HEX = "F7F8FA"
AZUL = RGBColor(0x00, 0x3D, 0xA5)
AZUL_ESC = RGBColor(0x00, 0x2D, 0x72)
TEXTO = RGBColor(0x1F, 0x1F, 0x1F)
CINZA = RGBColor(0x5C, 0x66, 0x70)
FONTE = "Calibri"


def shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run(par, text: str, *, bold: bool = False, italic: bool = False,
            size: int = 11, color: RGBColor = TEXTO,
            font: str = FONTE, mono: bool = False):
    r = par.add_run(text)
    r.font.name = (mono and "Consolas") or font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return r


INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"        # bold
    r"|(\*[^*]+\*)"             # italic
    r"|(`[^`]+`)"               # code
    r"|(\[[^\]]+\]\([^)]+\))",  # link [txt](url) — exibimos só o txt
    flags=re.UNICODE,
)


def add_inline(par, text: str, *, base_size: int = 11, base_color=TEXTO):
    """Renderiza markdown inline em runs do parágrafo."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            add_run(par, text[pos:m.start()], size=base_size, color=base_color)
        token = m.group(0)
        if token.startswith("**"):
            add_run(par, token[2:-2], bold=True, size=base_size, color=base_color)
        elif token.startswith("*"):
            add_run(par, token[1:-1], italic=True, size=base_size, color=base_color)
        elif token.startswith("`"):
            add_run(par, token[1:-1], size=base_size, color=base_color, mono=True)
        elif token.startswith("["):
            mtxt = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mtxt:
                add_run(par, mtxt.group(1), size=base_size, color=AZUL)
        pos = m.end()
    if pos < len(text):
        add_run(par, text[pos:], size=base_size, color=base_color)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        add_run(p, text, bold=True, size=20, color=AZUL_ESC)
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), AMARELO_HEX)
        bdr.append(bottom)
        pPr.append(bdr)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        add_run(p, text, bold=True, size=15, color=AZUL)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        add_run(p, text, bold=True, size=12, color=AZUL_ESC)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        add_run(p, text, bold=True, size=11, color=AZUL_ESC)
    return p


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    add_inline(p, text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    add_inline(p, text, base_color=CINZA)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), AMARELO_HEX)
    bdr.append(left)
    pPr.append(bdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str, *, ordered: bool = False, number: int | None = None):
    p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    add_inline(p, text)
    p.paragraph_format.space_after = Pt(2)


def parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    cols = max(len(header), max((len(r) for r in rows), default=0))
    if cols == 0:
        return
    header += [""] * (cols - len(header))
    for r in rows:
        r += [""] * (cols - len(r))

    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.autofit = True

    # Cabeçalho amarelo BB
    for j, txt in enumerate(header):
        cell = table.rows[0].cells[j]
        shade(cell, AMARELO_HEX)
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], txt, base_size=10)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = AZUL_ESC

    # Corpo
    for i, row in enumerate(rows, start=1):
        for j, txt in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], txt, base_size=10)
    # zebra striping suave
    for i in range(1, 1 + len(rows)):
        if i % 2 == 0:
            for cell in table.rows[i].cells:
                shade(cell, CINZA_FUNDO_HEX)


def render_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Linha em branco
        if not line.strip():
            i += 1
            continue

        # Separador horizontal --- ignora
        if re.match(r"^-{3,}$", line.strip()):
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # Tabela
        if line.lstrip().startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[i + 1]):
            header = parse_table_row(line)
            i += 2  # pula divisor
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Blockquote
        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            add_blockquote(doc, text)
            i += 1
            continue

        # Lista numerada
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            add_bullet(doc, m.group(1), ordered=True)
            i += 1
            continue

        # Lista com -
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            add_bullet(doc, m.group(1))
            i += 1
            continue

        # Parágrafo padrão
        add_paragraph(doc, line)
        i += 1


def build_docx(md_path: Path, out_path: Path, titulo: str, subtitulo: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # cabeçalho BB
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(h, titulo, bold=True, size=22, color=AZUL_ESC)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, subtitulo, bold=True, size=12, color=AZUL)
    doc.add_paragraph()

    md_text = md_path.read_text(encoding="utf-8")
    render_markdown(doc, md_text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK {out_path}")


def main() -> None:
    docs = ROOT / "docs"
    build_docx(
        docs / "ENTREGA_TIME1_MODELOS.md",
        docs / "ENTREGA_TIME1_MODELOS.docx",
        titulo="Entrega MVP — Time 1 (Modelos Analíticos)",
        subtitulo="HyperCopa DISEC 2026 · CESUP-Contratações · 07/05/2026",
    )
    build_docx(
        docs / "ENTREGA_TIME2_AGENTE.md",
        docs / "ENTREGA_TIME2_AGENTE.docx",
        titulo="Entrega MVP — Time 2 (Agente IA + RPA)",
        subtitulo="HyperCopa DISEC 2026 · CESUP-Contratações · 07/05/2026",
    )


if __name__ == "__main__":
    main()
