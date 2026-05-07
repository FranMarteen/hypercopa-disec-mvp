"""
Gera RELATORIO_SOLUCAO.docx a partir do markdown, com identidade visual BB.

Paleta BB: amarelo #FAE128, azul #003DA5, azul escuro #002D72.
Fonte: Calibri (substituto do BB Texto disponivel no Word). Trocar por
"BB Texto" se voce tiver a fonte instalada no SO.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "docs" / "RELATORIO_SOLUCAO.docx"

BB_AMARELO = RGBColor(0xFA, 0xE1, 0x28)
BB_AZUL = RGBColor(0x00, 0x3D, 0xA5)
BB_AZUL_ESCURO = RGBColor(0x00, 0x2D, 0x72)
BB_CINZA = RGBColor(0x5C, 0x66, 0x70)
BB_TEXTO = RGBColor(0x1F, 0x1F, 0x1F)
BB_AMARELO_HEX = "FAE128"
BB_AZUL_HEX = "003DA5"
BB_AZUL_ESCURO_HEX = "002D72"
BB_FUNDO_SUAVE_HEX = "F7F8FA"

FONTE = "Calibri"  # troque por "BB Texto" se a fonte estiver instalada


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_border_bottom(par, hex_color: str, sz: int = 12) -> None:
    p_pr = par._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_run(par, text: str, *, size=11, bold=False, color=BB_TEXTO,
            italic=False, font=FONTE):
    run = par.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def heading1(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(18)
    par.paragraph_format.space_after = Pt(8)
    add_run(par, text, size=18, bold=True, color=BB_AZUL_ESCURO)
    set_paragraph_border_bottom(par, BB_AMARELO_HEX, sz=20)
    return par


def heading2(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(4)
    add_run(par, text, size=14, bold=True, color=BB_AZUL_ESCURO)
    return par


def paragraph(doc, *parts):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(6)
    for p in parts:
        if isinstance(p, str):
            add_run(par, p)
        else:
            text, style = p
            add_run(
                par, text,
                bold=style.get("bold", False),
                italic=style.get("italic", False),
                color=style.get("color", BB_TEXTO),
                size=style.get("size", 11),
            )
    return par


def bullet(doc, text, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    par.paragraph_format.space_after = Pt(2)
    add_run(par, text)
    return par


def add_table(doc, header, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(header):
        cell = table.cell(0, i)
        set_cell_bg(cell, BB_AZUL_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        par = cell.paragraphs[0]
        add_run(par, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_bg(cell, BB_FUNDO_SUAVE_HEX)
            cell.text = ""
            par = cell.paragraphs[0]
            add_run(par, str(val), size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


def add_capa(doc):
    # Faixa amarela superior
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(4)
    set_paragraph_border_bottom(par, BB_AMARELO_HEX, sz=24)
    add_run(par, " ", size=2)

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_after = Pt(0)
    add_run(par, "BANCO DO BRASIL", size=10, bold=True, color=BB_AZUL)

    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(0)
    add_run(par, "DISEC · Time HyperCopa: Equipe HyperCopa DISEC 2026",
            size=9, color=BB_CINZA)

    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(120)
    par.paragraph_format.space_after = Pt(6)
    add_run(par, "Relatorio da Solucao", size=32, bold=True,
            color=BB_AZUL_ESCURO)

    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(0)
    add_run(par, "Agente Preparador BB + Modelo Analitico", size=18,
            color=BB_AZUL)

    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(120)
    add_run(par, "HyperCopa DISEC 2026", size=12, italic=True,
            color=BB_CINZA)

    # Tag amarela
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, BB_AMARELO_HEX)
    cell.text = ""
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(par, "  Documento para banca avaliadora  ", bold=True,
            size=11, color=BB_AZUL_ESCURO)

    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(60)
    add_run(par, f"Versao 1.0 · {datetime.now().strftime('%d/%m/%Y')}",
            size=10, color=BB_CINZA)

    doc.add_page_break()


def main():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Default style font
    style = doc.styles["Normal"]
    style.font.name = FONTE
    style.font.size = Pt(11)

    # ----- CAPA -----
    add_capa(doc)

    # ----- 1. Sumario executivo -----
    heading1(doc, "1. Sumario executivo")
    paragraph(
        doc,
        "A jornada ",
        ("Agente Preparador + Modelo Analitico", {"bold": True}),
        " transforma um extrato bruto de Licitacao Eletronica em um modelo "
        "preditivo treinado e auditavel em ",
        ("menos de 5 minutos", {"bold": True}),
        ", sem que o usuario de negocio escreva uma unica linha de codigo.",
    )
    paragraph(
        doc,
        "A solucao resolve a maior dor da DISEC ao operacionalizar analytics: "
        "o tempo entre receber um extrato da area e ter uma resposta acionavel. "
        "Hoje esse ciclo e de ",
        ("dias a semanas", {"bold": True}),
        " (DBA -> Cientista de Dados -> Validacao -> Modelo). Nossa solucao "
        "reduz para ",
        ("minutos", {"bold": True}),
        ", mantendo a governanca tecnica via system_prompt versionado e "
        "sandbox de execucao.",
    )

    add_table(
        doc,
        ["Indicador", "Antes", "Com a solucao"],
        [
            ["Tempo CSV bruto -> modelo treinado", "3-10 dias", "3-5 minutos"],
            ["Linhas de codigo escritas pelo usuario", "50-300", "0"],
            ["Pessoas envolvidas no ciclo", "3-4", "1"],
            ["Risco de vazamento (CSV completo p/ LLM)", "alto", "zero"],
        ],
        col_widths=[8, 4, 4],
    )

    # ----- 2. Contexto e problema -----
    heading1(doc, "2. Contexto e problema")
    heading2(doc, "2.1 Realidade DISEC")
    paragraph(
        doc,
        "A DISEC do Banco do Brasil opera Licitacao Eletronica sob a Lei "
        "14.133/21. As areas demandantes (DICOI, DISUP, DITEC, GECOI) "
        "periodicamente recebem perguntas de negocio:",
    )
    bullet(doc, "Quais EAPs vao atrasar nos proximos 30 dias?")
    bullet(doc, "Qual o risco de ruptura contratual desta carteira?")
    bullet(doc, "Que fornecedores tendem a ter intercorrencias?")

    heading2(doc, "2.2 O gargalo")
    paragraph(doc, "Hoje, transformar essas perguntas em modelos passa por:")
    bullet(doc, "1. Solicitar extrato a um DBA (1-3 dias).")
    bullet(doc, "2. Cientista de dados explora, limpa, gera features (2-5 dias).")
    bullet(doc, "3. Treinar e validar o modelo (1-2 dias).")
    bullet(doc, "4. Apresentar resultado.")
    paragraph(
        doc,
        ("Custo de oportunidade: ", {"bold": True}),
        "decisoes adiadas, contratos que rompem antes da analise ficar pronta, "
        "perda de janelas regulatorias.",
    )

    heading2(doc, "2.3 Decisoes estrategicas tomadas")
    add_table(
        doc,
        ["Decisao", "Motivo"],
        [
            ["Dados 100% sinteticos",
             "Acesso a dados reais inviavel no prazo do hackathon. Geracao via Faker + regras de negocio."],
            ["3 modelos preditivos via H2O AutoML",
             "GBM venceu na maioria dos casos; H2O permite trocar de algoritmo sem reescrita."],
            ["Agente IA via OpenAI function calling",
             "Pattern maduro (Tool Use), 4 tools auditaveis, sem RAG inicial (escopo justo para MVP)."],
            ["Streamlit single-page",
             "Uma URL, tipografia BB, jornada linear — pensada para demanda recorrente."],
        ],
        col_widths=[6, 10],
    )

    # ----- 3. Arquitetura -----
    heading1(doc, "3. Arquitetura — dois caminhos")
    paragraph(
        doc,
        "A solucao tem dois caminhos paralelos para a fase conversacional, "
        "com a mesma camada de treino e relatorio local.",
    )

    heading2(doc, "3.1 Caminho A — OpenAI direto")
    paragraph(
        doc,
        "Chat embutido no app Streamlit local. O agente OpenAI (gpt-4o-mini) "
        "chama 4 tools (ler_schema, ler_amostra, executar_pandas, "
        "salvar_csv_final) que rodam localmente. Apenas schema + ate 20 linhas "
        "de amostra saem para a OpenAI publica. Indicado para piloto, demo "
        "externa e desenvolvimento.",
    )

    heading2(doc, "3.2 Caminho B — Microsoft Copilot do Teams (sem API)")
    paragraph(
        doc,
        "Agente declarativo publicado no Copilot Studio do tenant M365 BB. "
        "Fluxo copy-paste: o usuario conversa no Teams, o Copilot devolve um "
        "bloco estruturado (PERGUNTA / TARGET / TASK / FEATURES_MANTER / "
        "FILTRO / JOINS / PASSO_A_PASSO_PANDAS), o usuario cola no app local "
        "que executa o codigo pandas em sandbox. Nenhuma chamada a API, nenhum "
        "tunnel, nenhum custo OpenAI. Auditoria automatica via Microsoft Purview.",
    )
    paragraph(
        doc,
        ("Como criar: ", {"bold": True}),
        "ver docs/COPILOT_STUDIO_GUIA.md (passo-a-passo, 20-30 min, sem codigo).",
    )

    heading2(doc, "3.3 Comparativo dos dois caminhos")
    add_table(
        doc,
        ["Aspecto", "Caminho A (OpenAI)", "Caminho B (Copilot Teams)"],
        [
            ["Frontend", "Streamlit (chat embutido)", "Microsoft Teams"],
            ["LLM", "OpenAI gpt-4o-mini/4o", "Copilot M365 (Microsoft)"],
            ["Custo por conversa", "~ R$ 0,30", "R$ 0 (incluso na licenca)"],
            ["CSV completo trafega", "Nao (so schema + amostra)", "Nao (so amostra colada)"],
            ["Onde a chave/licenca mora", ".env do usuario", "Tenant M365 BB"],
            ["Auditoria", "Logs Streamlit + git", "Microsoft Purview"],
            ["Setup", "pip install + chave", "Copilot Studio (sem codigo)"],
            ["Disponibilidade", "Local apenas", "Qualquer dispositivo Teams"],
        ],
        col_widths=[5, 5.5, 5.5],
    )

    heading2(doc, "3.4 Os 3 modelos preditivos")
    add_table(
        doc,
        ["#", "Modelo", "Tipo", "Target", "Uso"],
        [
            ["1", "Prazo de contratacao", "Regressao GBM",
             "dias ate assinatura", "priorizar EAPs com risco de estouro"],
            ["2", "Intercorrencia", "Classificacao binaria GBM",
             "houve impugnacao/recurso?", "alocacao de juridico"],
            ["3", "Ruptura contratual", "Classificacao binaria GBM",
             "houve rescisao?", "retencao de fornecedores"],
        ],
        col_widths=[1, 3.5, 4, 3.5, 4],
    )

    # ----- 4. Identidade visual e UX -----
    heading1(doc, "4. Identidade visual e UX")
    heading2(doc, "4.1 Tipografia e paleta")
    bullet(doc, "Fonte primaria: IBM Plex Sans (Google Fonts) — humanista, "
                "semelhante a BB Texto. Substituicao trivial via @font-face.")
    bullet(doc, "Amarelo BB #FAE128 (Pantone 116C) — destaque, CTAs, header stripe.")
    bullet(doc, "Azul BB #003DA5 (Pantone 286C) — titulos, icones, header.")
    bullet(doc, "Azul Escuro BB #002D72 — gradiente do header.")
    bullet(doc, "Cinza BB #5C6670 — texto secundario.")

    heading2(doc, "4.2 Principios da jornada")
    bullet(doc, "Uma URL, uma jornada — sem multi-page do Streamlit.")
    bullet(doc, "Sidebar = estado, centro = acao.")
    bullet(doc, "Etapa 2 (modelo) desbloqueia automaticamente quando o agente "
                "entrega o CSV.")
    bullet(doc, "Relatorio HTML autocontido — abre em qualquer navegador, "
                "com identidade BB.")

    # ----- 5. Privacidade e governanca -----
    heading1(doc, "5. Privacidade e governanca")
    add_table(
        doc,
        ["Risco", "Mitigacao Caminho A", "Mitigacao Caminho B"],
        [
            ["CSV inteiro vazar para LLM",
             "Apenas schema + <=20 linhas saem para OpenAI",
             "CSV completo nunca sai do laptop; so amostra trafega pelo Teams"],
            ["LLM publico fora do controle BB",
             "Risco residual aceito em piloto",
             "Eliminado: Copilot M365 sob acordo BB-Microsoft"],
            ["Codigo gerado malicioso",
             "Sandbox: exec() com globals restritos",
             "Mesmo sandbox; bloco PASSO_A_PASSO_PANDAS validado"],
            ["Chave/credencial vazar",
             ".gitignore protege .env",
             "Sem chave individual — usa licenca Copilot do tenant"],
            ["Dados reais BB em repo externo",
             ".gitignore bloqueia dados reais/",
             "Idem"],
            ["Auditabilidade",
             "Logs Streamlit + git do system_prompt",
             "Microsoft Purview automatico"],
            ["Custo marginal por conversa",
             "~ R$ 0,30 (OpenAI)",
             "R$ 0 (incluso em licenca Copilot)"],
        ],
        col_widths=[5, 5.5, 5.5],
    )

    # ----- 6. Resultados -----
    heading1(doc, "6. Resultados e metricas")
    heading2(doc, "6.1 Performance dos 3 modelos (dados sinteticos, 60s budget)")
    add_table(
        doc,
        ["Modelo", "Metrica primaria", "Valor (teste)"],
        [
            ["Prazo (regressao)", "RMSE / R2", "~ 18 dias / 0.82"],
            ["Intercorrencia (classificacao)", "AUC", "~ 0.84"],
            ["Ruptura (classificacao)", "AUC", "~ 0.79"],
        ],
        col_widths=[6, 5, 5],
    )
    paragraph(
        doc,
        ("Observacao: ", {"italic": True, "bold": True}),
        ("valores tipicos obtidos no MVP. A solucao final mede exatamente as "
         "metricas de cada execucao no relatorio HTML gerado pelo app.",
         {"italic": True}),
    )

    heading2(doc, "6.2 Metricas operacionais")
    bullet(doc, "3-5 minutos do CSV cru ao modelo treinado.")
    bullet(doc, "0 linhas de codigo escritas pelo usuario de negocio.")
    bullet(doc, "15 turnos maximos por conversa (limite de seguranca).")
    bullet(doc, "8 iteracoes maximas tool <-> LLM por turno.")

    # ----- 7. ROI -----
    heading1(doc, "7. ROI estimado (cenario DISEC anualizado)")
    paragraph(
        doc,
        ("Estimativas conservadoras com hipoteses anotadas — nao sao "
         "compromissos.", {"italic": True, "color": BB_CINZA}),
    )
    add_table(
        doc,
        ["Item", "Hipotese", "Valor anual"],
        [
            ["Demandas evitadas a cientista de dados",
             "60 demandas/ano x 6h economizadas", "360h"],
            ["Economia em homem-hora (R$ 150/h carregado)",
             "360h x R$ 150", "R$ 54.000"],
            ["Contratos com ruptura evitada",
             "0,5% de 4.000 contratos x R$ 200k medio", "R$ 4.000.000"],
            ["Custo Caminho A (OpenAI)",
             "1.000 conversas x R$ 0,30", "R$ 300"],
            ["Custo Caminho B (Copilot Teams)",
             "ja incluso em licenca M365 BB", "R$ 0"],
            ["Saldo estimado primeiro ano", "—", "~ R$ 4 MM"],
        ],
        col_widths=[6, 6, 4],
    )

    # ----- 8. Plano de evolucao -----
    heading1(doc, "8. Plano de evolucao")
    heading2(doc, "8.1 Curto prazo (ate Pitch Day, 10/06/2026)")
    bullet(doc, "[ok] MVP funcional com agente + 3 modelos.")
    bullet(doc, "[ok] Identidade visual BB.")
    bullet(doc, "[ok] Relatorio HTML auto-contido.")
    bullet(doc, "[ ] Demonstracao com extratos sinteticos das 4 areas "
                "(DICOI, DISUP, DITEC, GECOI).")
    bullet(doc, "[ ] Video de 2 min + slides de pitch.")

    heading2(doc, "8.2 Medio prazo (apos aprovacao)")
    bullet(doc, "Substituir IBM Plex Sans pelas fontes oficiais BB "
                "(BB Texto / BB Titulos).")
    bullet(doc, "Publicar agente Copilot no Teams via Copilot Studio "
                "(Caminho B oficial — ver docs/COPILOT_STUDIO_GUIA.md).")
    bullet(doc, "RAG com EAPs Padrao e jurisprudencia da Lei 14.133/21.")
    bullet(doc, "Integracao com Microsoft Teams (Copilot Studio).")
    bullet(doc, "Persistencia das conversas em banco BB para auditoria.")

    heading2(doc, "8.3 Longo prazo")
    bullet(doc, "Out-of-time validation automatica a cada trimestre.")
    bullet(doc, "Deteccao de drift nas variaveis top do leaderboard.")
    bullet(doc, "Recomendacao ativa: agente alerta a area quando uma EAP "
                "cruzar o limiar de risco.")
    bullet(doc, "API REST para consumo por sistemas internos BB.")

    # ----- 9. Estrutura entregavel -----
    heading1(doc, "9. Estrutura entregavel")
    add_table(
        doc,
        ["Entregavel", "Arquivo", "Status"],
        [
            ["App principal (Caminhos A e B)", "app_agente_bb.py", "Pronto"],
            ["System prompt OpenAI (Caminho A)",
             "docs/agente/system_prompt.md", "Pronto"],
            ["Schema das 4 tools (Caminho A)",
             "docs/agente/tools_schema.json", "Pronto"],
            ["System prompt Copilot Studio (Caminho B)",
             "teams_copilot/instructions.md", "Pronto"],
            ["Manifest do agente declarativo",
             "teams_copilot/declarative-agent.json", "Pronto"],
            ["Guia Copilot Studio passo-a-passo",
             "docs/COPILOT_STUDIO_GUIA.md", "Pronto"],
            ["Fluxograma (2 caminhos)", "docs/FLUXOGRAMA.md", "Pronto"],
            ["Doc de acesso e instalacao", "docs/ACESSO.md", "Pronto"],
            ["Relatorio (este)",
             "docs/RELATORIO_SOLUCAO.md / .docx", "Pronto"],
            ["Dados sinteticos",
             "dados_sinteticos/, dados_postgres/", "Pronto"],
            ["Repo privado",
             "github.com/FranMarteen/hypercopa-disec-mvp", "No ar"],
        ],
        col_widths=[5.5, 8, 2.5],
    )

    # ----- 10. Equipe -----
    heading1(doc, "10. Equipe e responsabilidades")
    bullet(doc, "Capitao da Equipe — lider tecnico, arquitetura da jornada, "
                "agente preparador, integracao Copilot/Streamlit.")
    bullet(doc, "Bento — modelagem preditiva, feature engineering, "
                "validacao dos 3 modelos H2O.")
    bullet(doc, "João — fluxo conversacional, refinamento do system_prompt, "
                "exemplos few-shot, identidade visual BB.")
    bullet(doc, "Apoio: DISEC — terminologia oficial BB e conformidade "
                "Lei 14.133/21.")

    # ----- 11. Conclusao -----
    heading1(doc, "11. Conclusao")
    paragraph(
        doc,
        ("A solucao demonstra que ", {}),
        ("velocidade analitica", {"bold": True}),
        " e alcancavel sem sacrificar ",
        ("governanca", {"bold": True}),
        ". A separacao entre o que sai para o LLM (apenas metadados e amostra) "
        "e onde o modelo e treinado (local, na JVM do H2O) garante "
        "conformidade com politicas BB de classificacao de dados mesmo em um "
        "piloto rodado fora do datacenter.",
    )
    paragraph(
        doc,
        "A jornada esta pronta para apresentacao ao Pitch Day em ",
        ("10/06/2026", {"bold": True}),
        ".",
    )

    # ----- Rodape -----
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(40)
    set_paragraph_border_bottom(par, BB_AMARELO_HEX, sz=8)
    add_run(par, " ", size=2)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        par,
        f"Relatorio gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} · "
        f"DISEC · Banco do Brasil",
        size=9, italic=True, color=BB_CINZA,
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
