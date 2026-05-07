"""Gera docs/MVP_CANVAS.docx com identidade visual BB."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "docs" / "MVP_CANVAS.docx"

AMARELO_HEX = "FAE128"
AZUL_HEX = "003DA5"
AZUL_ESCURO_HEX = "002D72"
CINZA_FUNDO_HEX = "F7F8FA"
AZUL = RGBColor(0x00, 0x3D, 0xA5)
AZUL_ESC = RGBColor(0x00, 0x2D, 0x72)
TEXTO = RGBColor(0x1F, 0x1F, 0x1F)
CINZA = RGBColor(0x5C, 0x66, 0x70)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONTE = "Calibri"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def run(par, text, *, bold=False, italic=False, size=11, color=TEXTO, font=FONTE):
    r = par.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return r


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run(p, text, bold=True, size=18, color=AZUL_ESC)
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), AMARELO_HEX)
        bdr.append(bottom)
        pPr.append(bdr)
    elif level == 2:
        run(p, text, bold=True, size=13, color=AZUL)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(doc, text, *, italic=False, size=11):
    p = doc.add_paragraph()
    run(p, text, italic=italic, size=size)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run(p, text)
    p.paragraph_format.space_after = Pt(2)


def canvas_row(table, label_text, sublabel, content_paragraphs):
    """Adiciona linha do canvas: rótulo amarelo à esquerda, conteúdo à direita."""
    row = table.add_row()
    cell_left, cell_right = row.cells
    shade(cell_left, AMARELO_HEX)
    shade(cell_right, "FFFFFF")

    p = cell_left.paragraphs[0]
    run(p, label_text, bold=True, size=12, color=AZUL_ESC)
    p2 = cell_left.add_paragraph()
    run(p2, sublabel, italic=True, size=9, color=AZUL_ESC)

    cell_right.paragraphs[0].text = ""
    for kind, text in content_paragraphs:
        if kind == "p":
            pp = cell_right.add_paragraph()
            run(pp, text)
            pp.paragraph_format.space_after = Pt(4)
        elif kind == "b":
            pp = cell_right.add_paragraph(style="List Bullet")
            run(pp, text)
        elif kind == "h":
            pp = cell_right.add_paragraph()
            run(pp, text, bold=True, size=11, color=AZUL)


def main():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Cabeçalho
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(h, "MVP Canvas — HyperCopa DISEC 2026", bold=True, size=22, color=AZUL_ESC)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(sub, "Agente Predfy + Modelo Analítico (H2O AutoML)", bold=True, size=13, color=AZUL)

    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(eq, "Equipe: Equipe HyperCopa DISEC 2026  |  Banco do Brasil · DISEC · CESUP-Contratações  |  07/05/2026",
        size=10, color=CINZA)

    doc.add_paragraph()

    # Tabela canvas
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11.5)

    # 1. Personas
    canvas_row(
        table,
        "1. Personas segmentadas",
        "Para que área foi destinado esse MVP?",
        [
            ("h", "Categoria do desafio"),
            ("p", "Inteligência Acionável para Compras e Fornecedores."),
            ("h", "Personas atendidas (DISEC – CESUP-Contratações)"),
            ("b", "Demandantes de áreas-cliente (DICOI, DISEC, DITEC, GECOI) — gestores que recebem perguntas analíticas e hoje dependem de fila DBA + Cientista de Dados."),
            ("b", "Líderes de contratação da DISEC — quem prioriza carteira sob a Lei 13.303/16."),
            ("b", "Cientistas de dados internos — deixam de fazer ETL repetitivo e focam em modelos de maior valor."),
        ],
    )

    # 2. Proposta
    canvas_row(
        table,
        "2. Proposta do MVP",
        "Problema de negócio. Que ações foram simplificadas/melhoradas?",
        [
            ("h", "Problema"),
            ("p", "O ciclo de uma pergunta de negócio (\"qual a chance de ruptura desta carteira?\") até a resposta acionável dura 3 a 10 dias, passando por DBA → Cientista de Dados → Validação. Decisões são adiadas; contratos rompem antes da análise ficar pronta."),
            ("h", "O que foi simplificado/melhorado"),
            ("b", "Tempo do extrato bruto ao modelo: 3-10 dias → 3-5 minutos."),
            ("b", "Linhas de código que o usuário precisa escrever: 50-300 → zero."),
            ("b", "Pessoas no ciclo: 3-4 → 1."),
            ("b", "Caminho A (OpenAI direto): só schema + 20 linhas de amostra trafegam para o LLM. Caminho B (Copilot Teams): CSV completo nunca sai do laptop."),
            ("h", "Como funciona"),
            ("p", "Usuário sobe CSV no app BB; um Agente IA (OpenAI ou Microsoft Copilot do Teams) entende a pergunta, propõe target/features e gera CSV final via sandbox pandas. H2O AutoML treina 3 modelos preditivos em até 5 min. Copilot do Teams traduz o relatório técnico em recomendações de negócio."),
        ],
    )

    # 3. Jornadas
    canvas_row(
        table,
        "3. Jornadas",
        "Jornadas atendidas. Indicadores estratégicos do Plano DISEC 2026: Integração para a Ação.",
        [
            ("h", "Jornadas DISEC atendidas"),
            ("b", "Priorização de carteira de EAPs — prever risco de atraso em contratações em andamento."),
            ("b", "Gestão de fornecedores — antecipar intercorrências (impugnação, recurso) antes da homologação."),
            ("b", "Mitigação de ruptura contratual — identificar contratos com alta probabilidade de rescisão."),
            ("b", "Atendimento a demanda recorrente das áreas-cliente, sem fila de cientista de dados."),
            ("h", "Indicadores estratégicos contemplados (DISEC 2026 – Integração para a Ação)"),
            ("b", "Tempo de resposta a demandas analíticas (3-10 dias → 3-5 min)."),
            ("b", "Aderência à governança de dados BB (Caminho B mantém CSV completo no laptop; só amostra trafega no tenant M365 sob Microsoft Purview)."),
            ("b", "Capilaridade analítica — qualquer demandante usa, sem precisar saber Python."),
            ("b", "Conformidade Lei 13.303/16 — terminologia oficial (Licitação Eletrônica, EAPs Padrão) embutida no system_prompt versionado em git."),
        ],
    )

    # 4. Tecnologias
    canvas_row(
        table,
        "4. Tecnologias envolvidas",
        "Quais tecnologias foram utilizadas? Justifique a escolha.",
        [
            ("h", "Modelo Analítico"),
            ("b", "H2O AutoML 3.46 (GBM, GLM, XGBoost, RF, DRF) — roda em JVM local; dados nunca saem do laptop. AutoML elimina escolha manual de algoritmo. Leaderboard auditável. Open-source."),
            ("h", "Agente de IA"),
            ("b", "Caminho A — OpenAI gpt-4o-mini com Tool Use (4 tools: ler_schema, ler_amostra, executar_pandas, salvar_csv_final). Sandbox exec() restrito. ~R$ 0,30/conversa. Apropriado para demo."),
            ("b", "Caminho B (produção BB) — Microsoft Copilot do Teams via Copilot Studio (declarativo, sem código). Trafega no tenant M365 BB sob acordo Microsoft↔BB. Auditoria automática Microsoft Purview. R$ 0 marginal (incluso na licença Copilot)."),
            ("h", "RPA / Automação"),
            ("b", "Copilot do Teams interpreta JSON do relatório H2O em recomendações executivas (RPA conversacional). Sandbox pandas substitui scripts ad-hoc do cientista de dados."),
            ("h", "Frontend / UX"),
            ("b", "Streamlit single-page com identidade visual BB (paleta #FAE128 / #003DA5, IBM Plex Sans como proxy de BB Texto). Jornada linear; sidebar = estado, centro = ação."),
            ("h", "Dados"),
            ("b", "Sintéticos (Faker + regras Lei 13.303/16). Acesso a dados reais inviável no prazo. Inclui ciclo pós-contrato (aditivos, rescisão, atrasos)."),
        ],
    )

    # 5. Resultados
    canvas_row(
        table,
        "5. Resultados Obtidos",
        "Que aprendizados ou resultados foram obtidos?",
        [
            ("h", "Resultados técnicos do MVP (07/05/2026)"),
            ("b", "Modelo 1 — Prazo (regressão GBM): RMSE ~18 dias / R² ~0,82."),
            ("b", "Modelo 2 — Intercorrência (classif. binária GBM): AUC ~0,84."),
            ("b", "Modelo 3 — Ruptura contratual (classif. binária GBM): AUC ~0,79."),
            ("b", "App entregável funcional (app_agente_bb.py) com identidade visual BB e jornada de 3 etapas."),
            ("b", "Dois caminhos validados: OpenAI (A) e Microsoft Copilot Teams (B). B é o de produção BB."),
            ("b", "Repositório privado no GitHub com .gitignore protegendo .env e dados reais/."),
            ("h", "Aprendizados"),
            ("b", "Velocidade analítica e governança não são trade-offs: separar o que sai para o LLM (metadados+amostra) de onde o modelo é treinado (JVM H2O local)."),
            ("b", "AutoML > escolher algoritmo manual: GBM venceu na maioria dos targets, mas a leaderboard expõe a escolha quando não."),
            ("b", "Caminho B é o de produção: elimina chave individual, custo OpenAI e dependência de LLM público."),
            ("b", "Sandbox restrito é não-negociável: código gerado por LLM precisa de globals limitados (sem os, sem subprocess, sem rede)."),
        ],
    )

    # 6. Métricas
    canvas_row(
        table,
        "6. Métricas para validação",
        "Indicadores afetados (tempo, custo, qualidade, risco). Como evidenciamos. Ganhos na entrega.",
        [
            ("h", "Indicadores afetados"),
            ("b", "Tempo CSV bruto → modelo treinado: 3-10 dias → 3-5 min (cronometrado no vídeo)."),
            ("b", "Linhas de código escritas pelo demandante: 50-300 → 0."),
            ("b", "Pessoas no ciclo: 3-4 → 1."),
            ("b", "Vazamento de CSV para LLM: alto → 0% (Caminho A só amostra; Caminho B nem amostra completa sai)."),
            ("b", "Custo marginal por conversa: R$ 0,30 (A) / R$ 0 (B, incluso na licença Copilot)."),
            ("b", "AUC dos modelos de risco: 0,84 (intercorrência) / 0,79 (ruptura)."),
            ("b", "Auditabilidade: system_prompt versionado em git + leaderboard H2O + logs Streamlit."),
            ("h", "Ganhos obtidos na data da entrega (07/05/2026)"),
            ("b", "App funcional e entregável (Caminhos A e B)."),
            ("b", "3 modelos validados com performance acima do baseline (AUC > 0,75)."),
            ("b", "Documentação completa: fluxograma, acesso, relatório, guia Copilot Studio."),
            ("b", "Vídeo de demonstração da jornada ponta-a-ponta."),
        ],
    )

    # 7. Escalabilidade
    canvas_row(
        table,
        "7. Escalabilidade",
        "Outras áreas beneficiadas. Previsão de ganhos (horas, financeiro).",
        [
            ("h", "Outras áreas BB que podem ser beneficiadas (mesmo padrão de jornada)"),
            ("b", "DIRAO / DIROP — risco de inadimplência por carteira."),
            ("b", "DICOI / DITEC / DISEC / GECOI — atraso em projetos de TI / obras (já contemplada)."),
            ("b", "DICAR / DIPES — turnover por agência ou função."),
            ("b", "CRGOV / DIRIS — risco regulatório por contrato (com RAG de cláusulas)."),
            ("h", "Previsão de ganhos (cenário DISEC anualizado, hipóteses conservadoras)"),
            ("b", "60 demandas/ano × 6h economizadas = 360 horas."),
            ("b", "360h × R$ 150/h carregada = R$ 54.000 em homem-hora/ano."),
            ("b", "0,5% de 4.000 contratos × R$ 200 mil = R$ 4.000.000/ano de ruptura evitada."),
            ("b", "Custo marginal Caminho A: R$ 300/ano; Caminho B: R$ 0 (incluso na licença M365)."),
            ("b", "Saldo estimado primeiro ano DISEC: ~ R$ 4 milhões."),
            ("h", "Caminho de escala"),
            ("b", "Curto prazo (até 10/06/2026): publicar agente Copilot no Teams (Copilot Studio); pilotar com extratos das 4 áreas DISEC."),
            ("b", "Médio prazo: RAG com EAPs Padrão e jurisprudência da Lei 13.303/16; persistência das conversas em banco BB; substituir IBM Plex Sans pelas fontes oficiais BB."),
            ("b", "Longo prazo: out-of-time validation trimestral; detecção de drift; agente proativo (alerta quando EAP cruza limiar de risco); API REST para sistemas internos BB."),
        ],
    )

    # Rodapé
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(foot, "Repositório: github.com/FranMarteen/hypercopa-disec-mvp (privado)  ·  ",
        size=9, color=CINZA)
    run(foot, "Artefatos: app_agente_bb.py, docs/FLUXOGRAMA.md, docs/ACESSO.md, docs/RELATORIO_SOLUCAO.md",
        size=9, color=CINZA)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Canvas salvo em: {OUT}")


if __name__ == "__main__":
    main()
